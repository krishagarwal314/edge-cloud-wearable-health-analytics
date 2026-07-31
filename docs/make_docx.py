"""Build the project documentation as a .docx for upload to Google Docs.

Google Docs converts .docx with full fidelity — headings, tables and images all survive,
which Markdown import does not guarantee. The content here is the same material that lives
in the repository's Markdown files; this script assembles it into a single submission
document.

    pip install python-docx
    python docs/make_docx.py

Output: docs/Project_Documentation.docx
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = pathlib.Path(__file__).resolve().parents[1]
FIG = REPO / "results" / "figures"
OUT = REPO / "docs" / "Project_Documentation.docx"

INK = RGBColor(0x1C, 0x20, 0x24)
ACCENT = RGBColor(0x25, 0x63, 0xA8)
MUTED = RGBColor(0x5B, 0x66, 0x72)

TITLE = "Cloud-Based Wearable Health Analytics Platform using Edge–Cloud Intelligence"
REPO_URL = "https://github.com/krishagarwal314/edge-cloud-wearable-health-analytics"

TEAM = [
    ("Krish Agarwal", "23BIT0427", "Cloud Infrastructure, DevOps, Backend & Data", "~45 %"),
    ("Monis Raza", "23BIT228", "AI/ML & Edge Intelligence", "~30 %"),
    ("Rudra Srivastav", "23BIT174", "Frontend, Visualisation & Documentation", "~25 %"),
]


# --------------------------------------------------------------------------- helpers
def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, colour in (("Heading 1", 17, ACCENT), ("Heading 2", 13.5, INK),
                               ("Heading 3", 11.5, INK)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True


def h(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text, bold=False, italic=False, size=None, align=None, colour=None,
         space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold, run.italic = bold, italic
    if size:
        run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def shade(cell, hex_colour="E8EEF6"):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    tcPr.append(el)


def table(doc, header, rows, widths=None, font=9.5, header_fill="D9E3F0"):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Word/Google Docs ignore cell widths unless the layout is explicitly fixed.
    t.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    t._tbl.tblPr.append(layout)
    hdr = t.rows[0].cells
    for i, txt in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(str(txt)); r.bold = True; r.font.size = Pt(font)
        shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val)); r.font.size = Pt(font)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
        for i, w in enumerate(widths):
            t.columns[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(doc, path: pathlib.Path, caption: str, width=6.4):
    if not path.exists():
        para(doc, f"[missing figure: {path.name}]", italic=True, colour=MUTED)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = para(doc, caption, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
             colour=MUTED, space_after=12)
    return c


def code_block(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return p


# --------------------------------------------------------------------------- build
def build() -> None:
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.85)

    # ---------------------------------------------------------------- title page
    for _ in range(3):
        doc.add_paragraph()
    para(doc, TITLE, bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, colour=ACCENT,
         space_after=10)
    para(doc, "Project Documentation — Planning & Design Phase", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, colour=MUTED, space_after=28)
    para(doc, "Course: Cloud Computing", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    table(doc, ["#", "Name", "Registration No.", "Area of Responsibility"],
          [[i + 1, n, r, a] for i, (n, r, a, _) in enumerate(TEAM)],
          widths=[0.4, 1.9, 1.5, 3.1], font=10.5)

    doc.add_paragraph()
    para(doc, "GitHub Repository", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=2)
    para(doc, REPO_URL, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, colour=ACCENT,
         space_after=18)
    para(doc, "Status: planning and design phase — code implementation not required at "
              "this stage.", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
         colour=MUTED)

    # ---------------------------------------------------------------- contents
    h(doc, "Table of Contents", 1, page_break=True)
    for i, s in enumerate([
        "Problem Statement", "Objectives", "Literature Survey", "Research Gap Analysis",
        "Proposed Architecture / Framework", "Technology Stack", "Dataset Details",
        "Repository Structure", "Work Distribution & Individual Contribution",
        "Work Completed So Far", "Limitations, Ethics & Future Work", "References"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{i}.  ").bold = True
        p.add_run(s)

    # ---------------------------------------------------------------- 1 problem
    h(doc, "1.  Problem Statement", 1, page_break=True)
    para(doc,
         "Consumer wearables — smartwatches, chest straps and pulse oximeters — generate a "
         "continuous, high-frequency stream of physiological data including heart rate, ECG, "
         "blood-oxygen saturation, skin temperature and accelerometry. In practice this data "
         "remains largely under-used, for three reasons.")
    bullets(doc, [
        ("Bandwidth, energy and cost.  ",
         "Continuously streaming raw high-rate signals such as 125 Hz ECG to the cloud for "
         "every user is expensive and drains device battery, and the overwhelming majority "
         "of those samples are clinically uninteresting."),
        ("Latency of critical alerts.  ",
         "A cardiac event detected only after a round trip to a remote data centre — or "
         "after an overnight batch job — is detected too late to be clinically useful."),
        ("Fragmentation.  ",
         "Vendor applications silo the data. Clinicians and researchers receive summary "
         "dashboards rather than queryable longitudinal records, and cannot substitute "
         "their own models."),
    ])
    h(doc, "Problem addressed by this project", 2)
    para(doc,
         "How can we build a low-cost, scalable health-analytics pipeline that detects "
         "physiological anomalies in near-real-time, without transmitting every raw sample "
         "to the cloud, while still preserving enough data for longitudinal analysis and "
         "model retraining?", italic=True)
    para(doc,
         "Our answer is an edge–cloud split-intelligence architecture. A lightweight "
         "quantised anomaly detector runs on the wearable gateway and forwards only "
         "(a) periodic compact summaries and (b) full-fidelity signal windows surrounding "
         "suspicious events. The cloud tier — built entirely from AWS Free Tier serverless "
         "components — performs confirmation inference, persistence, alerting and "
         "visualisation.")

    # ---------------------------------------------------------------- 2 objectives
    h(doc, "2.  Objectives", 1)
    h(doc, "Primary objectives", 2)
    table(doc, ["ID", "Objective"], [
        ["O1", "Design and implement an end-to-end edge-to-cloud telemetry pipeline for "
               "wearable vital signs using AWS IoT Core, AWS Lambda and Amazon DynamoDB."],
        ["O2", "Train a lightweight anomaly-detection model (1-D convolutional autoencoder "
               "over ECG/HR windows), quantise it to TensorFlow Lite under 500 KB, and "
               "execute it on the edge gateway."],
        ["O3", "Implement a two-stage inference cascade — cheap edge screening followed by "
               "higher-capacity cloud confirmation — and quantify the bandwidth saved "
               "relative to a naive stream-everything baseline."],
        ["O4", "Deploy the entire cloud tier within AWS Free Tier limits, targeting a $0 "
               "monthly bill, with all infrastructure defined as code (CloudFormation/SAM)."],
        ["O5", "Deliver a responsive web dashboard for live vitals, historical trends and an "
               "alert timeline, secured with Amazon Cognito."],
    ], widths=[0.5, 6.4])
    h(doc, "Secondary objectives", 2)
    table(doc, ["ID", "Objective"], [
        ["O6", "Provide an event-driven alerting path (SNS → email/SMS) with under 5 s "
               "end-to-end latency from anomalous sample to notification."],
        ["O7", "Benchmark the system: ingestion throughput, end-to-end latency, model "
               "accuracy (precision, recall, F1) and cost-per-1000-devices projection."],
        ["O8", "Demonstrate elasticity by replaying a synthetic multi-device load and "
               "showing automatic Lambda concurrency scaling."],
    ], widths=[0.5, 6.4])

    # ---------------------------------------------------------------- 3 lit survey
    h(doc, "3.  Literature Survey", 1, page_break=True)
    h(doc, "3.1  Survey method", 2)
    bullets(doc, [
        ("Databases searched:  ", "IEEE Xplore, ACM Digital Library, ScienceDirect, "
                                  "PubMed, arXiv."),
        ("Inclusion criteria:  ", "English, peer-reviewed or high-citation preprint, "
                                  "published 2015–2025, reporting either a system "
                                  "architecture or a quantitative evaluation."),
        ("Exclusion criteria:  ", "purely clinical studies with no computational "
                                  "contribution; papers without evaluation."),
    ])

    h(doc, "3.2  Theme A — IoT / cloud architectures for health monitoring", 2)
    table(doc, ["Ref", "Work", "Contribution", "Limitation"], [
        ["[1]", "Islam et al., IEEE Access, 2015",
         "Foundational taxonomy of IoT healthcare across sensing, network, service and "
         "application layers.",
         "Pre-dates practical edge ML; cloud-centric; no cost analysis."],
        ["[2]", "Baker, Xiang & Atkinson, IEEE Access, 2017",
         "Generic end-to-end IoT healthcare framework; identifies energy, security and "
         "interoperability as open issues.",
         "Framework is conceptual; never deployed or measured."],
        ["[3]", "Rahmani et al., FGCS, 2018",
         "Smart e-health gateway performing local processing, compression and embedded "
         "storage at the fog layer.",
         "Dedicated on-premise hardware; no public-cloud integration or elasticity."],
    ], widths=[0.45, 1.5, 2.6, 2.35], font=9)

    h(doc, "3.3  Theme B — Edge and fog intelligence", 2)
    table(doc, ["Ref", "Work", "Contribution", "Limitation"], [
        ["[4]", "Shi et al., IEEE IoT Journal, 2016",
         "Defines the edge-computing paradigm and its latency, bandwidth and privacy "
         "motivations.", "Vision paper; no health-specific instantiation."],
        ["[5]", "Satyanarayanan, IEEE Computer, 2017",
         "Cloudlet model; argues for tiered compute placement.",
         "Conceptual; no split-inference design."],
        ["[6]", "Kang et al., ASPLOS, 2017 (Neurosurgeon)",
         "Automatically partitions a DNN between mobile and cloud at the optimal layer.",
         "Partitions a single network; assumes continuous connectivity; not an "
         "anomaly-triggered cascade."],
        ["[7]", "Teerapittayanon et al., ICDCS, 2017 (DDNN)",
         "Distributed DNN with early-exit branches: easy samples exit at the edge, hard "
         "ones escalate.",
         "Evaluated on vision, not physiological time series; no cost analysis."],
    ], widths=[0.45, 1.5, 2.6, 2.35], font=9)

    h(doc, "3.4  Theme C — Machine learning for physiological anomaly detection", 2)
    table(doc, ["Ref", "Work", "Contribution", "Limitation"], [
        ["[8]", "Moody & Mark, IEEE EMB Mag., 2001",
         "The MIT-BIH Arrhythmia Database — the reference annotated ECG corpus.",
         "Dataset paper; few subjects; clinical rather than wearable recording conditions."],
        ["[9]", "Hannun et al., Nature Medicine, 2019",
         "34-layer CNN matching cardiologist performance on 12 rhythm classes from "
         "single-lead ambulatory ECG.",
         "Far too large for a Raspberry Pi; inference assumed server-side."],
        ["[10]", "Chauhan & Vig, IEEE DSAA, 2015",
         "LSTM predictive model; anomalies flagged from the prediction-error distribution "
         "— semi-supervised, needing only normal data.",
         "Recurrent model expensive on-device; no system integration."],
        ["[11]", "Malhotra et al., ICML AD Workshop, 2016",
         "Reconstruction-error anomaly detection for multivariate sensor streams — the "
         "template for our autoencoder.",
         "Not quantised or deployed; offline evaluation only."],
        ["[12]", "Kiranyaz et al., IEEE TBME, 2016",
         "Compact 1-D CNN for real-time, patient-specific ECG classification.",
         "Requires per-patient labelled data; no cloud tier."],
    ], widths=[0.45, 1.5, 2.6, 2.35], font=9)

    h(doc, "3.5  Theme D — Serverless and cost-efficient cloud backends", 2)
    table(doc, ["Ref", "Work", "Contribution", "Limitation"], [
        ["[13]", "Jonas et al., Berkeley View, 2019",
         "Defines serverless computing, its economics and its limitations (state, cold "
         "starts, data locality).", "General-purpose; no IoT or health case study."],
        ["[14]", "Baldini et al., 2017",
         "Survey of FaaS platforms and design patterns.",
         "Predates most managed IoT-to-FaaS integrations."],
        ["[15]", "Aslanpour et al., ACSW, 2021",
         "Argues for FaaS abstractions spanning edge and cloud.",
         "Vision-level; no concrete healthcare pipeline with measured cost."],
    ], widths=[0.45, 1.5, 2.6, 2.35], font=9)

    h(doc, "3.6  Comparative summary", 2)
    table(doc, ["Aspect", "Cloud-only [1–3]", "Edge/fog [3–7]", "Deep ECG [8–12]",
                "This project"], [
        ["Where inference runs", "Cloud", "Edge only", "Offline / server",
         "Edge screen + cloud confirm"],
        ["Uplink volume", "Full raw stream", "Reduced, fixed policy", "Not applicable",
         "Adaptive, anomaly-triggered"],
        ["Model size constraint", "None", "Considered", "Ignored",
         "INT8 TFLite < 500 KB, measured"],
        ["Backend model", "VMs / containers", "On-prem gateway", "Not applicable",
         "Fully serverless (FaaS)"],
        ["Deployment cost reported", "Rarely", "Rarely", "Never",
         "Free Tier budget + projection"],
        ["Reproducible IaC", "No", "No", "No", "Yes — one-command deploy"],
        ["End-to-end latency measured", "Sometimes", "Sometimes", "No",
         "Yes, p95 target < 5 s"],
    ], widths=[1.45, 1.3, 1.3, 1.25, 1.6], font=8.5)

    # ---------------------------------------------------------------- 4 gap
    h(doc, "4.  Research Gap Analysis", 1, page_break=True)
    h(doc, "4.1  Identified gaps", 2)
    for title, body in [
        ("Gap 1 — Cascaded inference is unvalidated for physiological streams",
         "Early-exit and model-partitioning schemes [6][7] are demonstrated on image "
         "classification, where a sample is a discrete, independent input. Physiological "
         "data is a continuous, non-stationary, subject-specific stream in which the "
         "events of interest are rare (class imbalance frequently exceeds 100:1) and "
         "context-dependent — a heart rate of 160 bpm is normal while running and alarming "
         "at rest. No reviewed work evaluates an edge-screen to cloud-confirm cascade on "
         "wearable vital signs with an explicit false-negative budget."),
        ("Gap 2 — Bandwidth and energy savings are asserted rather than measured",
         "Fog-computing healthcare papers [3] claim reduced transmission but rarely publish "
         "a byte-level comparison against a stream-everything baseline under identical "
         "signal conditions, and almost never relate it to a monetary or quota cost."),
        ("Gap 3 — Deep ECG models exist in isolation from deployable systems",
         "State-of-the-art ECG models [9] are large — tens of megabytes, millions of "
         "parameters — and are evaluated offline on curated datasets. The literature does "
         "not carry them through quantisation, latency measurement on real edge hardware, "
         "and integration into a production data path."),
        ("Gap 4 — Health-IoT backends are rarely serverless, and cost is rarely reported",
         "Reviewed systems use virtual machines, containers or on-premise gateways. The "
         "serverless literature [13][14] analyses economics generically but not for a "
         "continuous-telemetry health workload. We found no health-IoT paper publishing an "
         "actual monthly bill or a per-device cost model."),
        ("Gap 5 — Reproducibility",
         "Architectures in this space are published as block diagrams. Almost none ship "
         "infrastructure-as-code, so a reader cannot re-create the reported results."),
        ("Gap 6 — Resilience of the edge link is under-treated",
         "Wearable gateways operate on intermittent connectivity, yet evaluations assume a "
         "stable uplink and data loss during disconnection is seldom quantified."),
    ]:
        h(doc, title, 3)
        para(doc, body)

    h(doc, "4.2  Gap-to-contribution matrix", 2)
    table(doc, ["Gap", "Our contribution", "How it is evaluated"], [
        ["G1", "Anomaly-triggered edge-to-cloud escalation with activity-aware thresholds",
         "Recall/precision versus uplink-volume operating curve"],
        ["G2", "Byte and message instrumentation against a stream-everything baseline",
         "Measured percentage reduction, reported in results/benchmarks/"],
        ["G3", "INT8 TFLite pipeline plus an on-device latency harness",
         "Accuracy delta and milliseconds per window on a Raspberry Pi 4"],
        ["G4", "Fully serverless tier with a published Free Tier budget",
         "AWS Cost Explorer bill plus a beyond-Free-Tier projection model"],
        ["G5", "One-command infrastructure-as-code and a hardware-free simulator",
         "Clean-account deployment from scratch, timed"],
        ["G6", "Prioritised store-and-forward buffer",
         "Chaos test: 30 minutes offline, measuring data loss and replay latency"],
    ], widths=[0.5, 3.2, 3.2], font=9)

    h(doc, "4.3  Research questions", 2)
    bullets(doc, [
        ("RQ1.  ", "For 10-second wearable vital-sign windows, what edge-threshold setting "
                   "minimises uplink volume subject to a false-negative rate below 5 %?"),
        ("RQ2.  ", "How much accuracy is lost by INT8 quantisation of the edge autoencoder, "
                   "and is that loss recovered by the cloud confirmation stage?"),
        ("RQ3.  ", "What is the p95 end-to-end latency from an anomalous sample to a "
                   "delivered notification in a purely serverless AWS pipeline, and what "
                   "dominates it?"),
        ("RQ4.  ", "At what device count does the architecture exit the AWS Free Tier, and "
                   "what is the marginal cost per device per month beyond that point?"),
        ("RQ5.  ", "Does incorporating accelerometer-derived activity context measurably "
                   "reduce motion-artefact false positives?"),
    ])

    h(doc, "4.4  Novelty statement", 2)
    para(doc,
         "Existing wearable-health platforms either stream all raw physiological data to the "
         "cloud — which is expensive, energy-hungry and slow to alert — or run detection "
         "entirely on the edge, which caps model capacity. Cascaded inference resolves this "
         "tension for computer vision, but has not been validated on continuous, highly "
         "imbalanced, context-dependent physiological streams. This project implements and "
         "empirically evaluates such a cascade end to end: a sub-500 KB quantised "
         "autoencoder screens 10-second windows on the device and escalates only suspicious "
         "ones to a higher-capacity confirmation model running in a fully serverless AWS "
         "backend. The contribution is threefold: an activity-aware escalation policy with a "
         "measured recall-versus-bandwidth operating curve; an end-to-end latency and "
         "accuracy evaluation of the deployed cascade rather than the model alone; and a "
         "fully reproducible, infrastructure-as-code deployment with a published cost model "
         "demonstrating zero-cost operation within the AWS Free Tier at demonstration scale.",
         italic=True)

    # ---------------------------------------------------------------- 5 architecture
    h(doc, "5.  Proposed Architecture / Framework", 1)
    para(doc,
         "The platform is a three-tier, event-driven, serverless architecture with split "
         "inference across the edge and cloud tiers.")
    figure(doc, FIG / "architecture_diagram.png",
           "Figure 1 — Proposed system architecture. The edge tier screens every window and "
           "escalates only suspicious ones; the cloud tier confirms, persists, alerts and "
           "serves the dashboard.", width=6.9)

    h(doc, "5.1  Tier responsibilities", 2)
    table(doc, ["Tier", "Responsibility", "Deployment"], [
        ["Edge", "Acquisition, signal conditioning, first-pass anomaly screening, adaptive "
                 "uplink, offline buffering", "Raspberry Pi gateway or Python simulator"],
        ["Cloud", "Ingestion, confirmation inference, persistence, alerting, query API",
         "AWS managed services (Free Tier)"],
        ["Presentation", "Clinician and caregiver dashboard",
         "React SPA on Amazon S3 + CloudFront"],
    ], widths=[1.0, 3.6, 2.3], font=9.5)

    h(doc, "5.2  The uplink policy (core mechanism)", 2)
    para(doc,
         "Every 10-second window is scored on the device; the resulting reconstruction "
         "error determines how much data is transmitted. This single decision is what keeps "
         "the deployment inside the AWS IoT Core free-tier message quota.")
    table(doc, ["Reconstruction error", "Flag", "What is published", "Approx. size"], [
        ["e < τ_low", "normal", "Compact statistical summary only", "≈ 430 bytes"],
        ["τ_low ≤ e < τ_high", "suspect", "Summary + downsampled waveform", "≈ 2 KB"],
        ["e ≥ τ_high", "critical", "Summary + full waveform, QoS 1, priority", "≈ 3.7 KB"],
    ], widths=[1.7, 0.9, 3.2, 1.1], font=9.5)
    para(doc,
         "Thresholds are activity-scaled: motion inflates reconstruction error, so tolerance "
         "rises while the wearer is moving. They are also overridable per device at runtime "
         "through an MQTT configuration topic, so the operating point can be tuned without "
         "reflashing devices.")

    h(doc, "5.3  Normal and anomaly data flows", 2)
    bullets(doc, [
        ("Normal path:  ",
         "sensor → conditioning → windowing → TFLite inference → summary published over "
         "MQTT → IoT Rule → ingest_handler Lambda → DynamoDB item with a 30-day TTL."),
        ("Anomaly path:  ",
         "the full waveform is attached and published at QoS 1 → ingest_handler archives it "
         "to Amazon S3 and asynchronously invokes anomaly_processor → the full-precision "
         "model plus a rule-based cross-check either confirms the event (writing an Alert "
         "and invoking alert_dispatcher, which de-duplicates and publishes to SNS) or "
         "rejects it, recording a hard negative for the next retraining round."),
    ])

    h(doc, "5.4  Key design decisions", 2)
    table(doc, ["#", "Decision", "Rationale"], [
        ["ADR-1", "AWS Lambda for all compute, not EC2",
         "Lambda's 1 M requests and 400 000 GB-seconds per month are always free, whereas "
         "the EC2 allowance expires after 12 months and bills continuously. Matches a "
         "bursty, event-driven IoT workload."],
        ["ADR-2", "DynamoDB as the hot store, not RDS",
         "25 GB of storage is always free; the access pattern (device partition key, "
         "timestamp sort key) fits naturally. The RDS free tier expires after 12 months."],
        ["ADR-3", "AWS IoT Core as the ingestion front door",
         "Managed MQTT with per-device X.509 identity, TLS, QoS 1 delivery on a lossy link, "
         "and a Rules Engine that invokes Lambda directly — no broker to operate."],
        ["ADR-4", "Split inference (edge screen + cloud confirm)",
         "Cloud-only inference wastes bandwidth and the message quota; edge-only inference "
         "cannot host a high-capacity model. The cascade delivers edge latency with cloud "
         "accuracy."],
        ["ADR-5", "INT8 post-training quantisation to TFLite",
         "Roughly 4× smaller and 3× faster on an ARM CPU; the accuracy delta is measured "
         "and reported rather than assumed."],
        ["ADR-6", "S3 + CloudFront for the web app",
         "Cheapest and simplest; the Amplify Hosting free tier is limited to 12 months."],
        ["ADR-7", "30-day TTL on telemetry items",
         "Keeps DynamoDB inside the 25 GB always-free limit; older data is already archived "
         "in S3."],
        ["ADR-8", "No VPC for Lambda functions",
         "A NAT Gateway costs approximately $32 per month and is never free."],
    ], widths=[0.55, 1.85, 4.5], font=9)

    h(doc, "5.5  Non-functional requirements", 2)
    table(doc, ["Requirement", "Target", "Verification"], [
        ["End-to-end alert latency", "< 5 s (p95) from sample to SNS publish",
         "AWS X-Ray traces and timestamp deltas"],
        ["Edge inference latency", "< 50 ms per 10 s window on a Raspberry Pi 4",
         "On-device timing harness"],
        ["Ingestion throughput", "≥ 50 messages/second (50 simulated devices)",
         "Load test in tests/integration/"],
        ["Uplink reduction", "≥ 90 % versus stream-everything",
         "Byte counters instrumented in the simulator"],
        ["Offline resilience", "Buffer ≥ 6 hours of data",
         "Chaos test: disconnect the network and verify replay"],
        ["Monthly cost", "$0 within the Free Tier at demonstration scale",
         "AWS Cost Explorer"],
        ["Security", "TLS in transit, encryption at rest, least-privilege IAM, no secrets "
                     "in the repository", "IAM policy review and automated secret scanning"],
        ["Privacy", "No personally identifying information in telemetry; opaque device IDs",
         "Schema review"],
    ], widths=[1.5, 2.7, 2.7], font=9)

    # ---------------------------------------------------------------- 6 stack
    h(doc, "6.  Technology Stack", 1, page_break=True)
    table(doc, ["Layer", "Technology", "Notes"], [
        ["Edge device", "Raspberry Pi 4 / Pi Zero 2 W, or a pure-Python simulator",
         "The simulator lets the project run with zero hardware"],
        ["Edge runtime", "Python 3.11, paho-mqtt, awsiotsdk, NumPy, SciPy", ""],
        ["Edge inference", "TensorFlow Lite runtime, INT8-quantised 1-D conv autoencoder",
         "Target size below 500 KB"],
        ["Ingestion", "AWS IoT Core (MQTT 3.1.1 over TLS 1.2), IoT Rules Engine",
         "250 000 messages/month free"],
        ["Compute", "AWS Lambda (Python 3.11), 128–512 MB", "Four functions"],
        ["Hot datastore", "Amazon DynamoDB (on-demand, TTL enabled)",
         "Three tables and one GSI; 25 GB always free"],
        ["Cold datastore", "Amazon S3 with a Standard → Infrequent Access lifecycle",
         "Raw windows and model artefacts"],
        ["API", "Amazon API Gateway HTTP API", "JWT authorizer"],
        ["Authentication", "Amazon Cognito user pool", "10 000 MAUs always free"],
        ["Notifications", "Amazon SNS", "Email and optional SMS"],
        ["Observability", "Amazon CloudWatch Logs, Metrics and Alarms; AWS X-Ray",
         "7-day log retention to stay within the free tier"],
        ["Infrastructure as code", "AWS SAM / CloudFormation (YAML)", "One-command deploy"],
        ["CI/CD", "GitHub Actions → sam deploy",
         "OIDC role assumption; no long-lived access keys"],
        ["Model training", "Python 3.11, TensorFlow/Keras, scikit-learn, Jupyter",
         "Google Colab or SageMaker Studio Lab (free)"],
        ["Frontend", "React 18 + Vite, Recharts, AWS Amplify Auth", ""],
        ["Web hosting", "Amazon S3 static website + CloudFront", "1 TB egress free"],
        ["Testing", "pytest, moto (AWS mocks), Vitest", ""],
    ], widths=[1.35, 2.75, 2.8], font=9)

    h(doc, "6.1  Free Tier budget", 2)
    para(doc,
         "The entire cloud tier is designed to run at zero cost. The binding constraint is "
         "the AWS IoT Core message quota, and the edge summarisation policy is what makes "
         "the project fit within it: at one message per second per device, five devices "
         "would consume roughly 13 million messages per month — fifty times over the limit. "
         "At one summary every ten seconds, five devices consume approximately 130 000.")
    table(doc, ["Service", "Free Tier allowance", "Type", "Planned usage"], [
        ["AWS IoT Core", "250 000 messages/month", "12 months", "≈ 130 000 (≈ 48 % used)"],
        ["AWS Lambda", "1 M requests + 400 000 GB-s/month", "Always free",
         "≈ 200 000 invocations (< 5 %)"],
        ["DynamoDB", "25 GB storage, 25 WCU + 25 RCU", "Always free",
         "≈ 2 GB with a 30-day TTL"],
        ["Amazon S3", "5 GB, 20 000 GET, 2 000 PUT/month", "12 months", "≈ 1 GB"],
        ["API Gateway", "1 M requests/month", "12 months", "≈ 50 000"],
        ["Amazon Cognito", "10 000 monthly active users", "Always free", "< 10 users"],
        ["Amazon SNS", "1 000 email notifications/month", "Always free", "< 500 alerts"],
        ["CloudWatch", "5 GB log ingestion, 10 alarms", "Always free",
         "Constrained by 7-day retention"],
        ["CloudFront", "1 TB egress/month", "Always free", "< 1 GB"],
    ], widths=[1.3, 2.1, 1.0, 2.5], font=9)
    para(doc,
         "Services deliberately avoided because they lack a usable free tier: Amazon "
         "Timestream, Kinesis Data Streams, SageMaker real-time endpoints, Amazon RDS, "
         "NAT Gateway, Amazon MSK and OpenSearch.", size=10, colour=MUTED)

    # ---------------------------------------------------------------- 7 dataset
    h(doc, "7.  Dataset Details", 1, page_break=True)
    para(doc,
         "All datasets are publicly released and de-identified. This project collects no new "
         "human-subject data, and raw data is not committed to the repository.")
    table(doc, ["Dataset", "Source", "Contents", "Used for"], [
        ["MIT-BIH Arrhythmia (primary)", "PhysioNet",
         "48 half-hour two-channel ambulatory ECG records, 360 Hz, 47 subjects, "
         "beat-level cardiologist annotations (~110 000 beats)",
         "Training the autoencoder on normal beats; anomaly-detection evaluation"],
        ["PTB-XL", "PhysioNet",
         "21 837 clinical 12-lead ECGs, 10 s each, 100/500 Hz, multi-label diagnoses",
         "Generalisation check — does the model transfer?"],
        ["WESAD", "UCI ML Repository",
         "15 subjects; chest (ECG, EDA, TEMP, RESP, ACC) and wrist (BVP, EDA, TEMP, ACC)",
         "Realistic wearable-grade noisy signals; multimodal fusion"],
        ["MHEALTH", "UCI ML Repository",
         "10 subjects, 23 channels, 12 physical activities",
         "Activity context, to suppress motion-artefact false positives (RQ5)"],
        ["Synthetic stream", "Generated in-repo (edge/simulator/)",
         "Configurable physiologically-shaped ECG across six profiles",
         "Load testing, elasticity demonstration, live demo"],
    ], widths=[1.35, 1.15, 2.4, 2.0], font=9)

    h(doc, "7.1  Preprocessing contract", 2)
    para(doc, "Every dataset is normalised to a common window format before training.")
    table(doc, ["Property", "Value"], [
        ["Signal", "Single-lead ECG (lead II or nearest equivalent)"],
        ["Filter", "Butterworth bandpass, 0.5–40 Hz, zero-phase"],
        ["Resampling", "125 Hz"],
        ["Window", "10 seconds (1 250 samples), 50 % overlap"],
        ["Normalisation", "Per-record z-score, fitted on that record only"],
        ["Auxiliary features", "HR mean/min/max, SDNN, SpO₂, temperature, activity class"],
        ["Label", "0 = normal, 1 = anomalous (any non-normal annotated beat in the window)"],
        ["Split", "Subject-disjoint 70/15/15 — no subject appears in two splits"],
    ], widths=[1.6, 5.3], font=9.5)
    para(doc,
         "Subject-disjoint splitting is treated as non-negotiable: splitting by window would "
         "leak recordings between training and test sets and inflate every reported metric.",
         italic=True, size=10, colour=MUTED)

    # ---------------------------------------------------------------- 8 repo
    h(doc, "8.  Repository Structure", 1, page_break=True)
    para(doc, "Every folder contains a README describing its purpose, so no directory is "
              "empty and the intent of each area is documented.")
    code_block(doc, """edge-cloud-wearable-health-analytics/
├── README.md                     Project overview (title, team, problem, objectives,
│                                 architecture, tech stack, dataset details)
├── CONTRIBUTING.md               Team conventions, branching, hard rules
├── LICENSE                       MIT
├── docs/
│   ├── WORK_DISTRIBUTION.md      Per-member responsibilities, timeline, ledger
│   ├── architecture/             Design document, API contract, Mermaid diagrams
│   ├── literature-survey/        Survey, research gap, BibTeX references
│   ├── aws/                      Free Tier budget, IAM and security notes
│   └── setup/                    Deployment runbook and troubleshooting
├── edge/                         Edge tier: simulator, TFLite inference, uplink policy
├── backend/                      Four AWS Lambda functions, shared layer, API spec
├── ai-models/                    Notebooks, training scripts, model artefacts
├── frontend/                     React dashboard (components, pages, services)
├── database/                     DynamoDB schemas, access patterns, seed data
├── infrastructure/               CloudFormation/SAM templates and deploy scripts
├── dataset/                      Dataset acquisition and preprocessing documentation
├── results/                      Benchmarks, figures, evaluation reports
├── presentation/                 Slides, demo script, report
├── tests/                        Unit and integration tests
└── .github/workflows/            CI pipeline""")

    table(doc, ["Folder", "Purpose", "Owner"], [
        ["edge/", "Signal conditioning, on-device inference, uplink policy, offline buffer",
         "Monis Raza"],
        ["backend/", "Serverless compute: ingest, anomaly processing, alerting, REST API",
         "Krish Agarwal"],
        ["ai-models/", "Model development from raw signal to a deployable TFLite artefact",
         "Monis Raza"],
        ["frontend/", "React dashboard for live vitals, trends and alerts",
         "Rudra Srivastav"],
        ["database/", "DynamoDB single-table design, access patterns, seed data",
         "Krish Agarwal"],
        ["infrastructure/", "Infrastructure as code, deployment and teardown scripts",
         "Krish Agarwal"],
        ["dataset/", "Dataset acquisition, licensing, preprocessing contract",
         "Monis Raza"],
        ["docs/", "Architecture, literature survey, AWS notes, setup runbook",
         "Rudra Srivastav"],
        ["results/", "Benchmarks, figures and evaluation reports", "Rudra Srivastav"],
        ["presentation/", "Slide deck, demo script, final report", "Rudra Srivastav"],
        ["tests/", "Unit and integration test suites", "Krish Agarwal"],
    ], widths=[1.3, 4.1, 1.5], font=9)

    # ---------------------------------------------------------------- 9 work dist
    h(doc, "9.  Work Distribution & Individual Contribution", 1, page_break=True)
    table(doc, ["Name", "Registration No.", "Area of Responsibility", "Tasks", "Share"],
          [[n, r, a, t, s] for (n, r, a, s), t in
           zip(TEAM, ["16 (I-1…I-8, B-1…B-8)", "8 (M-1…M-8)", "8 (F-1…F-8)"])],
          widths=[1.5, 1.2, 2.45, 1.15, 0.6], font=9.5)

    h(doc, "Rationale for the distribution", 2)
    para(doc,
         "The project comprises four distinct technical areas but the team has three "
         "members, so Krish Agarwal covers two of them — the cloud infrastructure tier and "
         "the backend/data tier. These are the most tightly coupled parts of the system: "
         "the CloudFormation template, the Lambda functions and the DynamoDB schema all "
         "change together, so dividing them between two people would create constant merge "
         "friction on the same files. This gives Krish 16 of the 32 planned tasks, with "
         "Monis and Rudra owning 8 each.")

    h(doc, "9.1  Krish Agarwal (23BIT0427) — Cloud Infrastructure, DevOps, Backend & Data", 2)
    para(doc, "Infrastructure and DevOps", bold=True, size=10.5, space_after=3)
    table(doc, ["ID", "Task", "Deliverable"], [
        ["I-1", "AWS account setup, IAM users, MFA, budget and billing alarm at $1",
         "docs/setup/SETUP.md"],
        ["I-2", "Author the SAM/CloudFormation template for the full stack",
         "infrastructure/cloudformation/template.yaml"],
        ["I-3", "IoT Core thing type, policy and X.509 certificate provisioning",
         "infrastructure/scripts/provision_device.sh"],
        ["I-4", "Least-privilege IAM roles for each Lambda function", "docs/aws/IAM_NOTES.md"],
        ["I-5", "CI/CD with GitHub Actions and OIDC federation to AWS",
         ".github/workflows/deploy.yml"],
        ["I-6", "CloudWatch dashboards, log-retention policies and alarms",
         "Dashboard definition and screenshots"],
        ["I-7", "Free Tier quota tracking and weekly cost reporting",
         "docs/aws/FREE_TIER_BUDGET.md"],
        ["I-8", "S3 + CloudFront static hosting and cache invalidation",
         "Template section and deploy script"],
    ], widths=[0.5, 3.6, 2.8], font=9)
    para(doc, "Backend and data", bold=True, size=10.5, space_after=3)
    table(doc, ["ID", "Task", "Deliverable"], [
        ["B-1", "Define and validate the telemetry message schema",
         "backend/common/schema.py, API contract"],
        ["B-2", "ingest_handler Lambda: validate, normalise, fan out to DynamoDB and S3",
         "backend/lambdas/ingest_handler/"],
        ["B-3", "anomaly_processor Lambda: cloud-side confirmation inference",
         "backend/lambdas/anomaly_processor/"],
        ["B-4", "alert_dispatcher Lambda: de-duplication, debounce, SNS publish",
         "backend/lambdas/alert_dispatcher/"],
        ["B-5", "api_handler Lambda and API Gateway routes",
         "backend/lambdas/api_handler/, OpenAPI spec"],
        ["B-6", "DynamoDB single-table design, GSIs and TTL policy", "database/schemas/"],
        ["B-7", "Cognito user pool integration and JWT authorizer wiring", "backend/api/"],
        ["B-8", "Integration tests with moto; load-test harness", "tests/integration/"],
    ], widths=[0.5, 3.6, 2.8], font=9)
    para(doc, "Also responsible for end-to-end system integration and for keeping the "
              "deployed stack within the AWS Free Tier.", size=10, italic=True)

    h(doc, "9.2  Monis Raza (23BIT228) — AI/ML & Edge Intelligence", 2)
    table(doc, ["ID", "Task", "Deliverable"], [
        ["M-1", "Acquire and document MIT-BIH, PTB-XL and WESAD; write download scripts",
         "dataset/, ai-models/preprocessing/"],
        ["M-2", "Preprocessing pipeline: filtering, resampling, windowing, normalisation",
         "ai-models/preprocessing/"],
        ["M-3", "Baseline models (threshold rules, Isolation Forest) for comparison",
         "ai-models/notebooks/01_baselines.ipynb"],
        ["M-4", "1-D convolutional autoencoder; training and threshold selection",
         "ai-models/training/train_autoencoder.py"],
        ["M-5", "INT8 post-training quantisation to TFLite; verify accuracy delta < 2 %",
         "ai-models/models/"],
        ["M-6", "Edge inference runtime and store-and-forward MQTT publisher",
         "edge/edge_inference/, edge/simulator/"],
        ["M-7", "Cloud-side confirmation model packaged for AWS Lambda",
         "ai-models/models/cloud/"],
        ["M-8", "Model evaluation: precision, recall, F1, ROC-AUC, confusion matrices",
         "results/benchmarks/model_eval.md"],
    ], widths=[0.5, 3.6, 2.8], font=9)

    h(doc, "9.3  Rudra Srivastav (23BIT174) — Frontend, Visualisation & Documentation", 2)
    table(doc, ["ID", "Task", "Deliverable"], [
        ["F-1", "React + Vite application scaffold, routing and theming", "frontend/"],
        ["F-2", "Cognito-backed login and logout flow", "frontend/src/services/auth.js"],
        ["F-3", "Live vitals view with polling updates", "frontend/src/pages/LiveVitals.jsx"],
        ["F-4", "Historical trends with a time-range picker", "frontend/src/pages/Trends.jsx"],
        ["F-5", "Alert timeline and acknowledgement action", "frontend/src/pages/Alerts.jsx"],
        ["F-6", "Device management view", "frontend/src/pages/Devices.jsx"],
        ["F-7", "Literature survey and research-gap documents", "docs/literature-survey/"],
        ["F-8", "Architecture diagrams, final report, slide deck and demo script",
         "docs/architecture/, presentation/"],
    ], widths=[0.5, 3.6, 2.8], font=9)

    h(doc, "9.4  Twelve-week timeline", 2)
    table(doc, ["Week", "Milestone", "Owner"], [
        ["1", "Problem finalisation and literature survey", "Rudra (all contribute)"],
        ["2", "Architecture design, repository scaffold, work split", "Krish (all contribute)"],
        ["3", "AWS account, IAM, billing alarms, IoT Core skeleton", "Krish"],
        ["4", "Telemetry schema, ingest_handler and DynamoDB tables", "Krish"],
        ["5", "Dataset acquisition and preprocessing pipeline", "Monis"],
        ["6", "Edge simulator publishing to IoT Core end to end", "Monis + Krish"],
        ["7", "Autoencoder training and quantisation", "Monis"],
        ["8", "Anomaly processor and SNS alerting path", "Krish"],
        ["9", "REST API and Cognito authentication", "Krish + Rudra"],
        ["10", "Dashboard: live view, trends, alerts", "Rudra"],
        ["11", "Benchmarking, elasticity demonstration, cost analysis", "Krish + Monis"],
        ["12", "Report, slides, demo rehearsal, submission", "All"],
    ], widths=[0.6, 4.2, 2.1], font=9.5)

    # ---------------------------------------------------------------- 10 progress
    h(doc, "10.  Work Completed So Far", 1, page_break=True)
    para(doc,
         "Code implementation is not required at this stage. Beyond the planning "
         "deliverables, we have built a working edge simulator so that the design can be "
         "validated end to end before any AWS resources are provisioned. The figures below "
         "are genuine output of that code, not illustrations.")
    table(doc, ["Deliverable", "Status"], [
        ["Problem definition, objectives, literature survey, research gap", "Complete"],
        ["Proposed architecture, ADRs, non-functional requirements", "Complete"],
        ["MQTT telemetry contract and REST API contract", "Complete"],
        ["Repository scaffold with a README in every folder", "Complete"],
        ["Work distribution and 12-week timeline", "Complete"],
        ["AWS Free Tier budget and IAM security plan", "Complete"],
        ["Edge simulator: ECG generation, uplink policy, bandwidth instrumentation",
         "Working prototype"],
        ["Telemetry JSON Schema", "Complete"],
        ["CI pipeline (lint, smoke test, template validation, secret scan)", "Complete"],
        ["Cloud tier deployment (CloudFormation)", "Planned — weeks 3–4"],
        ["Model training and quantisation", "Planned — weeks 5–7"],
        ["Dashboard implementation", "Planned — week 10"],
    ], widths=[4.8, 2.1], font=9.5)

    figure(doc, FIG / "sim_cascade_timeline.png",
           "Figure 2 — Edge screening across a simulated 400-window run (≈ 67 minutes). "
           "Windows below τ_low upload only a compact summary; escalated windows upload a "
           "waveform. Shaded bands mark ground-truth abnormal episodes.")
    figure(doc, FIG / "sim_bandwidth.png",
           "Figure 3 — Measured uplink volume against a stream-everything baseline: an "
           "86.8 % reduction. The baseline is itself compressed, so the comparison is not "
           "favourable by construction.")
    figure(doc, FIG / "dashboard_mockup.png",
           "Figure 4 — Design mockup of the clinician dashboard. This is a proposed "
           "interface, not a screenshot; the frontend is scheduled for week 10.")

    para(doc,
         "Note on interpretation: the anomaly scorer behind Figures 2 and 3 is a documented "
         "placeholder heuristic, not the trained TFLite autoencoder, which is scheduled for "
         "weeks 5–7. These figures demonstrate that the pipeline and the uplink policy "
         "function end to end; they are not the final model-accuracy results, and the "
         "bandwidth figure will be re-measured once the trained model is in place.",
         italic=True, size=10, colour=MUTED)

    # ---------------------------------------------------------------- 11 limitations
    h(doc, "11.  Limitations, Ethics & Future Work", 1, page_break=True)
    h(doc, "11.1  Scope and limitations", 2)
    bullets(doc, [
        ("Not a medical device.  ", "This is an academic prototype with no clinical "
         "validation and no regulatory claim. It must not be used for diagnosis or "
         "treatment."),
        ("Datasets are clinical, not wearable-grade.  ", "MIT-BIH was recorded with "
         "clinical electrodes, whereas wrist-worn photoplethysmography is considerably "
         "noisier. We mitigate this with noise and motion-artefact augmentation, but it "
         "remains a threat to external validity."),
        ("Limited demographic coverage.  ", "MIT-BIH contains only 47 subjects."),
        ("Optimistic prevalence.  ", "Anomaly prevalence in curated datasets is far higher "
         "than in a free-living population, so reported precision is optimistic."),
        ("Single region, no failover.  ", "Cross-region traffic is billable and out of "
         "scope; a disaster-recovery strategy is documented instead."),
        ("Free Tier constrains scale.  ", "Sustained load tests are short and bounded so "
         "the account remains within quota."),
    ])
    h(doc, "11.2  Ethics and privacy", 2)
    bullets(doc, [
        "All datasets used are publicly released and de-identified; no new human-subject "
        "data is collected, so no institutional ethics approval is required.",
        "The deployed system stores no personally identifying information — device "
        "identifiers are opaque, and any mapping to a person lives outside the system.",
        "Privacy improves as a direct consequence of the architecture: raw waveforms are "
        "transmitted only around suspected events, not continuously.",
        "Data is encrypted in transit (TLS 1.2 with mutual X.509 authentication) and at "
        "rest, with least-privilege IAM roles for every component.",
    ])
    h(doc, "11.3  Future work", 2)
    bullets(doc, [
        "Federated learning, so per-device models improve without centralising raw signals.",
        "On-device incremental personalisation using the wearer's own baseline.",
        "Multimodal fusion of ECG with electrodermal activity and respiration.",
        "Multi-region deployment with automated failover.",
        "A clinical validation study, which would be a prerequisite for any real use.",
    ])

    # ---------------------------------------------------------------- 12 references
    h(doc, "12.  References", 1)
    refs = [
        "Islam, S. M. R., Kwak, D., Kabir, M. H., Hossain, M., & Kwak, K.-S. (2015). The "
        "Internet of Things for Health Care: A Comprehensive Survey. IEEE Access, 3, "
        "678–708.",
        "Baker, S. B., Xiang, W., & Atkinson, I. (2017). Internet of Things for Smart "
        "Healthcare: Technologies, Challenges and Opportunities. IEEE Access, 5, "
        "26521–26544.",
        "Rahmani, A. M., et al. (2018). Exploiting Smart e-Health Gateways at the Edge of "
        "Healthcare Internet-of-Things: A Fog Computing Approach. Future Generation "
        "Computer Systems, 78, 641–658.",
        "Shi, W., Cao, J., Zhang, Q., Li, Y., & Xu, L. (2016). Edge Computing: Vision and "
        "Challenges. IEEE Internet of Things Journal, 3(5), 637–646.",
        "Satyanarayanan, M. (2017). The Emergence of Edge Computing. Computer, 50(1), 30–39.",
        "Kang, Y., et al. (2017). Neurosurgeon: Collaborative Intelligence Between the Cloud "
        "and Mobile Edge. ASPLOS.",
        "Teerapittayanon, S., McDanel, B., & Kung, H. T. (2017). Distributed Deep Neural "
        "Networks over the Cloud, the Edge and End Devices. IEEE ICDCS.",
        "Moody, G. B., & Mark, R. G. (2001). The Impact of the MIT-BIH Arrhythmia Database. "
        "IEEE Engineering in Medicine and Biology Magazine, 20(3), 45–50.",
        "Hannun, A. Y., et al. (2019). Cardiologist-Level Arrhythmia Detection and "
        "Classification in Ambulatory Electrocardiograms Using a Deep Neural Network. "
        "Nature Medicine, 25, 65–69.",
        "Chauhan, S., & Vig, L. (2015). Anomaly Detection in ECG Time Signals via Deep Long "
        "Short-Term Memory Networks. IEEE DSAA.",
        "Malhotra, P., et al. (2016). LSTM-Based Encoder-Decoder for Multi-Sensor Anomaly "
        "Detection. ICML Anomaly Detection Workshop.",
        "Kiranyaz, S., Ince, T., & Gabbouj, M. (2016). Real-Time Patient-Specific ECG "
        "Classification by 1-D Convolutional Neural Networks. IEEE Transactions on "
        "Biomedical Engineering, 63(3), 664–675.",
        "Jonas, E., et al. (2019). Cloud Programming Simplified: A Berkeley View on "
        "Serverless Computing. UC Berkeley EECS Technical Report UCB/EECS-2019-3.",
        "Baldini, I., et al. (2017). Serverless Computing: Current Trends and Open Problems. "
        "Research Advances in Cloud Computing, Springer.",
        "Aslanpour, M. S., et al. (2021). Serverless Edge Computing: Vision and Challenges. "
        "Australasian Computer Science Week.",
        "Goldberger, A. L., et al. (2000). PhysioBank, PhysioToolkit, and PhysioNet. "
        "Circulation, 101(23), e215–e220.",
        "Wagner, P., et al. (2020). PTB-XL, a Large Publicly Available Electrocardiography "
        "Dataset. Scientific Data, 7(154).",
        "Schmidt, P., et al. (2018). Introducing WESAD, a Multimodal Dataset for Wearable "
        "Stress and Affect Detection. ICMI.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(f"[{i}]  ").bold = True
        run = p.add_run(r)
        run.font.size = Pt(9.5)

    doc.add_paragraph()
    para(doc, f"Repository: {REPO_URL}", size=10, colour=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(str(OUT))
    print(f"wrote {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    build()
